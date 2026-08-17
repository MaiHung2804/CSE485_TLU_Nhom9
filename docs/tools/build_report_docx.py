from __future__ import annotations

import math
from datetime import datetime
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from table_geometry import apply_table_geometry, column_widths_from_weights


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = BASE_DIR / "bao-cao-de-tai-dat-san-the-thao-format-moi.docx"
LEGACY_OUTPUT_PATH = BASE_DIR / "bao-cao-vong-1.docx"
ASSET_DIR = BASE_DIR / "generated" / "report-assets"

COLOR_NAVY = RGBColor(0x12, 0x2A, 0x43)
COLOR_BLUE = RGBColor(0x1F, 0x5A, 0x94)
COLOR_SOFT_BLUE = RGBColor(0xE9, 0xF1, 0xF8)
COLOR_SOFT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
COLOR_MUTED = RGBColor(0x5B, 0x66, 0x73)
COLOR_GREEN = RGBColor(0x1F, 0x7A, 0x5A)
COLOR_BORDER = "C9D4E1"
COLOR_IMAGE_BG = (247, 250, 252)
COLOR_IMAGE_NAVY = (18, 42, 67)
COLOR_IMAGE_BLUE = (31, 90, 148)
COLOR_IMAGE_LIGHT = (233, 241, 248)
COLOR_IMAGE_GREEN = (223, 243, 233)
COLOR_IMAGE_ORANGE = (252, 236, 214)
COLOR_IMAGE_PINK = (250, 229, 232)
COLOR_IMAGE_GRAY = (88, 101, 115)

REPORT_META = {
    "institution": "TRUONG DAI HOC THUY LOI (TLU)",
    "faculty": "KHOA CONG NGHE THONG TIN",
    "subject": "Bao cao phan tich, thiet ke va xay dung he thong",
    "title": "Xay dung he thong quan ly va dat san the thao tai TLU",
    "topic_group": "De tai 6 - Nhom 2 - Dich vu sinh vien va khuon vien",
    "phase": "Pham vi trien khai: Vong 1",
    "student_name": "[Dien ho ten]",
    "student_id": "[Dien ma sinh vien]",
    "class_name": "[Dien lop]",
    "lecturer": "[Dien ten giang vien huong dan]",
    "technology": "Laravel 12, PHP 8.2, Blade, Bootstrap 5, MySQL, phpMyAdmin",
    "architecture": "MVC",
    "date_label": "Thang 08 nam 2026",
}

INTRO_PARAGRAPHS = [
    (
        "Trong boi canh chuyen doi so trong giao duc, nhu cau tin hoc hoa cac dich vu ho tro sinh vien "
        "khong chi duoc dat ra doi voi cac nghiep vu hoc tap ma con mo rong sang nhom dich vu trong khuon vien. "
        "Voi cac hoat dong the chat, giao luu va ren luyen, viec quan ly san the thao neu thuc hien thu cong "
        "se de phat sinh trung lich, kho theo doi tinh trang san va kho tong hop bao cao su dung."
    ),
    (
        "Tu bai toan thuc te do, de tai \"Xay dung he thong quan ly va dat san the thao tai TLU\" duoc lua chon "
        "nham xay dung mot he thong web ho tro quan ly danh muc san, lich mo, tiep nhan yeu cau dat san, "
        "kiem tra xung dot, phan quyen nguoi dung va tong hop thong ke su dung. He thong huong toi hai doi tuong "
        "chinh la quan tri vien va sinh vien, trong do moi doi tuong duoc cap dung chuc nang theo vai tro."
    ),
    (
        "Bao cao nay trinh bay toan bo qua trinh phan tich nghiep vu, xac dinh yeu cau, thiet ke co so du lieu, "
        "xay dung kien truc MVC va danh gia ket qua dat duoc cua phien ban hien tai. Noi dung duoc viet theo huong "
        "co the nop truc tiep cho vong 1, dong thoi de san kha nang mo rong cho cac vong tiep theo neu can."
    ),
]

GOALS = [
    "So hoa quy trinh quan ly va dat san the thao trong khuon vien truong.",
    "Han che tinh trang trung lich giua cac yeu cau dat san cung ngay va cung ca gio.",
    "Cho phep theo doi tinh trang san, lich mo, booking, usage log va bao cao su dung.",
    "Phan quyen ro rang giua admin va sinh vien theo dung workflow nghiep vu.",
    "Dam bao du an co the cai dat lai nhanh bang migration, seeder va du lieu demo.",
]

SCOPES = [
    "Tap trung vao huong trien khai \"San the thao\" trong De tai 6 cua giang vien.",
    "Phien ban hien tai su dung 8 bang nghiep vu cot loi: users, sport_types, courts, time_slots, court_schedules, bookings, booking_details, usage_logs.",
    "He thong chua trien khai thanh toan, thong bao email va lich bao tri nang cao theo ca.",
    "Bao cao dat trong pham vi project ca nhan lam local, quan ly ma nguon bang Git va day len GitHub.",
]

METHODS = [
    "Khao sat bai toan va xac dinh quy trinh nghiep vu dat san theo vai tro nguoi dung.",
    "Phan tich tac nhan, use case, quy tac rang buoc va cac tinh huong xung dot can xu ly.",
    "Thiet ke co so du lieu theo quan he cha - con va mo hinh MVC cho ung dung Laravel.",
    "Xay dung migration, seeder, model, controller, blade view va bo phan xac thuc, phan quyen.",
    "Kiem thu bang du lieu demo, kiem tra dang nhap, booking, duyet, huy va thong ke.",
]

REPORT_STRUCTURE = [
    ("Chuong 1", "Tong quan de tai"),
    ("Chuong 2", "Phan tich yeu cau he thong"),
    ("Chuong 3", "Thiet ke he thong"),
    ("Chuong 4", "Xay dung va cai dat he thong"),
    ("Ket luan", "Tong ket cac ket qua dat duoc cua de tai"),
]

ROLE_ROWS = [
    ("Admin", "Dang nhap, CRUD danh muc, duyet booking, cap nhat usage log, xem bao cao", "Quan ly toan bo tai nguyen va workflow he thong"),
    ("Sinh vien", "Dang nhap, tao booking, xem, sua, huy booking cua chinh minh", "Su dung san the thao phu hop nhu cau hoc tap va sinh hoat"),
]

MAIN_FUNCTION_ROWS = [
    ("Dang nhap va phan quyen", "Xac thuc nguoi dung va hien thi dung pham vi chuc nang theo vai tro admin / sinh vien."),
    ("Quan ly danh muc san", "Quan ly mon the thao, san, ca gio va tinh trang san."),
    ("Quan ly lich mo", "Gan lich mo cho tung san theo thu trong tuan va tung ca gio."),
    ("Dat san", "Sinh vien tao phieu dat, nhap muc dich, so nguoi, ngay dat va ca gio."),
    ("Phe duyet va huy booking", "Admin duyet / tu choi booking; sinh vien duoc huy booking cua minh khi con hop le."),
    ("Usage log va bao cao", "Cap nhat ket qua su dung thuc te va tong hop dashboard, thong ke."),
]

FUNCTIONAL_REQUIREMENTS = [
    "Dang nhap va dang xuat he thong theo tai khoan da cap.",
    "Quan ly danh muc mon the thao, san, ca gio va lich mo cua tung san.",
    "Cho phep sinh vien tao phieu dat san voi thong tin muc dich, so nguoi, ngay dat va ca gio.",
    "Kiem tra xung dot booking theo san, ngay, ca gio va tinh trang hoat dong cua san.",
    "Cho phep admin duyet hoac tu choi booking, luu ly do tu choi khi can.",
    "Cho phep sinh vien xem, sua, huy booking cua chinh minh khi con hop le.",
    "Cap nhat usage log sau buoi dat san de dong bo trang thai booking.",
    "Tong hop dashboard va bao cao thong ke theo khoang thoi gian.",
]

NON_FUNCTIONAL_REQUIREMENTS = [
    "Giao dien de hieu, viet hoa tieng Viet va co bo cuc ro rang cho form, bang va dashboard.",
    "Du lieu duoc validate o backend; khong thao tac thay doi du lieu bang GET.",
    "Mat khau duoc ma hoa; route quan tri duoc bao ve boi middleware va policy.",
    "He thong co the cai dat lai bang migrate:fresh --seed de phuc vu demo nhanh.",
    "Cau truc MVC ro rang, View khong query database va Controller khong chua SQL dai.",
]

USE_CASES = [
    ("UC01", "Dang nhap he thong", "Admin, Sinh vien", "Nguoi dung xac thuc de truy cap he thong theo vai tro."),
    ("UC02", "Quan ly mon the thao", "Admin", "Them, sua, xoa danh muc mon the thao phuc vu khai bao san."),
    ("UC03", "Quan ly san", "Admin", "Them, sua, xoa san; cap nhat suc chua, vi tri va trang thai."),
    ("UC04", "Quan ly ca gio", "Admin", "Khai bao cac khung gio su dung cho dat san."),
    ("UC05", "Quan ly lich mo san", "Admin", "Gan lich mo theo thu trong tuan va tung ca gio cho moi san."),
    ("UC06", "Tao booking", "Sinh vien", "Nhap thong tin dat san va gui yeu cau den he thong."),
    ("UC07", "Kiem tra xung dot", "He thong", "Chan dat san neu trung lich, san bao tri hoac khong mo lich."),
    ("UC08", "Duyet hoac tu choi booking", "Admin", "Cap nhat trang thai booking va ly do tu choi neu co."),
    ("UC09", "Huy booking", "Sinh vien", "Sinh vien huy phieu dat cua chinh minh khi chua hoan tat."),
    ("UC10", "Cap nhat usage log", "Admin", "Ghi nhan used, no_show hoac cancelled sau buoi dat san."),
    ("UC11", "Xem bao cao", "Admin", "Tong hop thong ke booking, usage log, top san va top sinh vien."),
]

WORKFLOW_STEPS = [
    "Admin khai bao mon the thao, san, ca gio va lich mo.",
    "Sinh vien chon san, ngay dat, ca gio va nhap thong tin booking.",
    "He thong kiem tra tinh trang san, suc chua, lich mo va xung dot.",
    "Neu hop le, booking duoc tao o trang thai pending.",
    "Admin duyet hoac tu choi booking.",
    "Sau buoi dat san, admin cap nhat usage log.",
    "Dashboard va bao cao tong hop du lieu su dung san.",
]

BUSINESS_RULES = [
    ("BR01", "San dang bao tri khong duoc phep dat.", "Kiem tra courts.status = active truoc khi tao booking."),
    ("BR02", "So nguoi tham gia khong duoc vuot qua suc chua cua san.", "So sanh player_count voi courts.capacity."),
    ("BR03", "San chi duoc dat trong khung gio da mo.", "Doi chieu court_schedules theo court_id, day_of_week va time_slot_id."),
    ("BR04", "Khong duoc co hai booking giao nhau cung san, cung ngay va cung ca gio.", "Tra cuu booking_details ket hop bookings co status pending hoac approved."),
    ("BR05", "Sinh vien chi duoc thao tac tren phieu dat cua chinh minh.", "Policy so sanh booking.user_id voi user hien tai."),
    ("BR06", "Chi admin moi duoc CRUD danh muc va duyet booking.", "Route staff-only duoc bao ve boi middleware role:admin."),
    ("BR07", "Moi booking detail chi co toi da mot usage log.", "unique tren usage_logs.booking_detail_id."),
]

TECHNOLOGIES = [
    ("PHP 8.2", "Ngon ngu lap trinh backend", "Phu hop Laravel va de trien khai local tren XAMPP."),
    ("Laravel 12", "Framework MVC", "Ho tro route, middleware, Eloquent, validation va Blade."),
    ("Blade", "Template engine", "Xay dung giao dien dashboard, form CRUD va trang booking."),
    ("Bootstrap 5", "CSS framework", "Giup giao dien responsive va thong nhat bo cuc."),
    ("MySQL", "He quan tri CSDL", "Luu du lieu booking, danh muc, lich mo va usage log."),
    ("phpMyAdmin", "Quan ly CSDL", "Quan sat bang va du lieu local trong qua trinh demo."),
    ("Git va GitHub", "Quan ly ma nguon", "Luu vet cac moc hoan thien va day du an len kho ma nguon."),
]

ENVIRONMENT_ROWS = [
    ("He dieu hanh", "Windows", "Moi truong phat trien local tren may ca nhan."),
    ("PHP", "8.2.12", "Chay Laravel thong qua XAMPP / CLI."),
    ("Composer", "2.10.2", "Quan ly dependency PHP."),
    ("Node.js", "v24.16.0", "Build tai nguyen frontend bang Vite."),
    ("npm", "11.13.0", "Cai dat package frontend."),
    ("Web server", "Apache (XAMPP)", "Phuc vu ung dung local."),
    ("Database", "MySQL (XAMPP)", "Luu du lieu nghiep vu cua he thong."),
    ("Cong cu CSDL", "phpMyAdmin", "Quan sat schema va du lieu demo."),
    ("Terminal", "Git Bash", "Thuc hien cac lenh cai dat, git va artisan."),
]

TABLE_GROUPS = [
    ("Nhom tai khoan", "users", "Luu thong tin dang nhap, vai tro va lien ket nguoi tao booking."),
    ("Nhom danh muc", "sport_types, courts, time_slots", "Quan ly mon the thao, san va khung gio."),
    ("Nhom lich", "court_schedules", "Quy dinh san nao mo vao thu nao, ca gio nao."),
    ("Nhom giao dich", "bookings, booking_details", "Luu phieu dat va chi tiet ngay dat, san dat, ca gio."),
    ("Nhom theo doi", "usage_logs", "Ghi nhan ket qua su dung thuc te sau buoi dat san."),
]

MODULE_ROWS = [
    ("Module xac thuc", "Dang nhap, dang xuat va phan quyen theo vai tro."),
    ("Module danh muc", "Quan ly mon the thao, san, ca gio va trang thai san."),
    ("Module lich mo", "Khai bao lich mo cho tung san theo thu va khung gio."),
    ("Module booking", "Tao, xem, sua, huy, duyet va tu choi booking."),
    ("Module usage log", "Cap nhat used, no_show, cancelled cho booking detail."),
    ("Module bao cao", "Thong ke booking, top san, top sinh vien va xu huong su dung."),
]

RELATIONSHIPS = [
    ("users", "bookings", "1 - N", "Mot nguoi dung co the tao nhieu phieu dat."),
    ("users", "bookings (approved_by)", "1 - N", "Mot admin co the duyet nhieu booking."),
    ("sport_types", "courts", "1 - N", "Mot mon the thao co nhieu san."),
    ("courts", "court_schedules", "1 - N", "Mot san co nhieu lich mo theo thu va ca."),
    ("time_slots", "court_schedules", "1 - N", "Mot ca gio co the ap dung cho nhieu san."),
    ("bookings", "booking_details", "1 - N", "Mot phieu dat co the mo rong nhieu dong chi tiet."),
    ("courts", "booking_details", "1 - N", "Mot san xuat hien trong nhieu booking detail."),
    ("time_slots", "booking_details", "1 - N", "Mot ca gio xuat hien trong nhieu booking detail."),
    ("booking_details", "usage_logs", "1 - 1", "Moi dong chi tiet co toi da mot usage log."),
    ("users", "usage_logs", "1 - N", "Admin cap nhat usage log cho tung buoi su dung."),
]

PERMISSION_ROWS = [
    ("Dang nhap / dang xuat", "Co", "Co", "Tat ca nguoi dung da cap tai khoan deu su dung duoc."),
    ("CRUD mon the thao", "Co", "Khong", "Chi admin quan ly danh muc sport_types."),
    ("CRUD san", "Co", "Khong", "Chi admin quan ly courts."),
    ("CRUD ca gio va lich mo", "Co", "Khong", "Chi admin quan ly time_slots va court_schedules."),
    ("Tao booking", "Khong", "Co", "Sinh vien chu dong gui yeu cau dat san."),
    ("Xem tat ca booking", "Co", "Khong", "Admin theo doi toan bo workflow."),
    ("Xem booking cua minh", "Khong", "Co", "Policy gioi han theo user_id."),
    ("Sua / huy booking cua minh", "Khong", "Co", "Chi khi booking chua o trang thai cancelled/completed."),
    ("Duyet / tu choi booking", "Co", "Khong", "Admin xu ly workflow pending."),
    ("Cap nhat usage log", "Co", "Khong", "Admin ghi nhan ket qua su dung."),
    ("Xem dashboard va bao cao", "Co", "Co mot phan", "Sinh vien xem dashboard ca nhan; bao cao tong hop danh cho admin."),
]

PROJECT_STRUCTURE = [
    ("app/Http/Controllers", "Chua cac controller xu ly dang nhap, booking, san, lich, usage log va bao cao."),
    ("app/Models", "Chua Eloquent model va relationship giua cac bang."),
    ("app/Http/Middleware", "Chua middleware role:admin cho route quan tri."),
    ("app/Policies", "Chua BookingPolicy de kiem soat quyen so huu booking."),
    ("database/migrations", "Khai bao cau truc bang va rang buoc khi tai lap CSDL."),
    ("database/seeders", "Nap du lieu demo cho tai khoan, san, lich mo, booking va usage log."),
    ("resources/views", "Giao dien Blade cho dashboard, booking, danh muc va bao cao."),
    ("routes/web.php", "Khai bao route, nhom middleware guest/auth/admin."),
    ("docs", "Tai lieu bao cao, ERD, data dictionary va file docx nop bai."),
]

COMPLETED_FEATURES = [
    ("Xac thuc va phan quyen", "Dang nhap, dang xuat, middleware role:admin va BookingPolicy."),
    ("Quan ly tai nguyen", "CRUD sport_types, courts, time_slots, court_schedules."),
    ("Workflow booking", "Tao booking, kiem tra xung dot, xem, sua, huy, duyet va tu choi."),
    ("Usage log", "Ghi nhan used, no_show, cancelled va dong bo trang thai booking."),
    ("Dashboard va bao cao", "Thong ke tong quan theo vai tro, top san, top sinh vien va xu huong dat san."),
    ("Du lieu demo", "Seeder tai khoan, san, lich mo, booking va usage log de demo nhanh."),
]

DEMO_ACCOUNTS = [
    ("admin@campus.local", "password", "Admin", "Demo CRUD, duyet booking, usage log va bao cao."),
    ("student1@campus.local", "password", "Sinh vien", "Demo tao booking va chinh sua booking cua minh."),
    ("student2@campus.local", "password", "Sinh vien", "Demo policy tu choi thao tac booking cua nguoi khac."),
]

DEMO_STEPS = [
    ("Buoc 1", "Bat Apache va MySQL trong XAMPP", "Dam bao localhost va MySQL san sang."),
    ("Buoc 2", "Mo Git Bash tai thu muc project", "Thao tac thong nhat dung nhu moi truong lam bai."),
    ("Buoc 3", "Chay php artisan migrate:fresh --seed", "Tai lap CSDL va nap du lieu demo."),
    ("Buoc 4", "Neu can giao dien, chay npm install va npm run dev", "Khoi dong Vite khi muon sua CSS/JS."),
    ("Buoc 5", "Chay php artisan serve", "Mo he thong tai http://127.0.0.1:8000."),
    ("Buoc 6", "Dang nhap admin va thuc hien duyet booking", "Chung minh workflow quan tri."),
    ("Buoc 7", "Dang nhap sinh vien va thu mo route quan tri", "Chung minh middleware va policy hoat dong."),
]

SCORE_ROWS = [
    ("A1", "Phan tich nghiep vu, ERD, thiet ke CSDL", "Da co 8 bang nghiep vu, PK/FK, rang buoc unique, data dictionary va so do ERD tong quan.", "Dat"),
    ("A2", "Khoi tao va tai lap CSDL", "Da co migration, seeder va SQL backup; co the chay migrate:fresh --seed.", "Dat"),
    ("A3", "Du lieu mau, Seeder/Factory", "Da co tai khoan demo, san, lich mo, booking pending/approved/completed va usage log.", "Dat"),
    ("A4", "Model/DAO va Relationship", "Da su dung Eloquent relationship dung chieu, fillable va truy van co to chuc.", "Dat"),
    ("A5", "CRUD/Workflow Dashboard", "Da co CRUD danh muc, workflow booking, dashboard va bao cao tong hop.", "Dat"),
    ("A6", "Validation, toan ven va bao mat nen", "Da co backend validation, CSRF, method dung, middleware va policy.", "Dat"),
    ("A7", "Kien truc va chat luong ma", "Da tach controller, model, view ro rang theo MVC.", "Dat"),
    ("A8", "UI va trai nghiem su dung", "Da Viet hoa giao dien, bo cuc dashboard, bang, form, flash message va responsive co ban.", "Dat"),
    ("A9", "Git, tich hop va dong gop ca nhan", "Da day code len GitHub; nen tiep tuc tach them commit theo moi moc nang cap de dep hon.", "Can cai thien"),
    ("A10", "README, ERD, demo va van dap", "Da co README, tai khoan demo va bo khung bao cao; can dien thong tin ca nhan tren bia truoc khi nop.", "Gan dat"),
]

LIMITATIONS = [
    "Chua co thong bao email hoac he thong nhac lich sau khi booking duoc duyet.",
    "Chua trien khai gioi han so lan dat san theo tuan cho tung sinh vien.",
    "Chua tach lich bao tri thanh mot module rieng co khung thoi gian cu the.",
    "Chua xuat bao cao sang PDF/Excel ngay trong giao dien web.",
]

FUTURE_DIRECTIONS = [
    "Bo sung booking policy nang cao: gioi han so slot / sinh vien / tuan.",
    "Them module bao tri san theo khoang thoi gian de khoa san tam thoi.",
    "Tich hop email, toast notification hoac thong bao trong he thong.",
    "Bo sung xuat bao cao PDF/Excel va lich su thao tac audit log.",
    "Mo rong dat nhieu ca gio trong cung mot phieu booking neu de tai yeu cau.",
]

REFERENCES = [
    "Tai lieu huong dan mon hoc va de bai thi cuoi ky CSE485.",
    "Tai lieu chinh thuc Laravel 12.x Documentation.",
    "MySQL 8.0 Reference Manual.",
    "Tai lieu su dung XAMPP va phpMyAdmin de trien khai local.",
    "Kho ma nguon du an: MaiHung2804/CSE485_TLU_Nhom9.",
]

APPENDIX_COMMANDS = [
    "cd \"/d/CN WEB LT/CN WEB LT/CN WEB/CSE485_TLU_Nhom9\"",
    "copy .env.example .env",
    "php artisan key:generate",
    "php artisan migrate:fresh --seed",
    "npm install",
    "npm run dev",
    "php artisan serve",
    "git add .",
    "git commit -m \"Hoan thien bao cao va project dat san the thao\"",
    "git push origin main",
]

APPENDIX_QA = [
    ("Tinh huong 1", "Dang nhap bang admin@campus.local", "Truy cap dashboard admin va thay menu quan tri day du."),
    ("Tinh huong 2", "Dang nhap bang student1@campus.local", "Truy cap dashboard sinh vien va tao booking moi."),
    ("Tinh huong 3", "Chon san dang bao tri", "He thong tu choi dat san va thong bao rang buoc."),
    ("Tinh huong 4", "Tao booking trung san, trung ngay, trung ca", "He thong chan xung dot va khong luu du lieu sai."),
    ("Tinh huong 5", "Dang nhap student2 va mo booking cua student1", "Policy chan sua/huy booking khong thuoc so huu."),
    ("Tinh huong 6", "Admin duyet booking", "Trang thai booking chuyen tu pending sang approved."),
    ("Tinh huong 7", "Admin cap nhat usage log", "Booking dong bo sang completed hoac cancelled phu hop."),
]

DATA_DICTIONARY = [
    {
        "name": "users",
        "note": "Bang luu thong tin tai khoan dang nhap va vai tro nguoi dung.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma nguoi dung"),
            ("name", "varchar(255)", "not null", "Ho ten hien thi"),
            ("email", "varchar(255)", "unique, not null", "Email dang nhap"),
            ("password", "varchar(255)", "not null", "Mat khau da hash"),
            ("role", "varchar(20)", "default student, index", "Vai tro admin/student"),
            ("remember_token", "varchar(100)", "nullable", "Token ghi nho dang nhap"),
            ("created_at", "timestamp", "nullable", "Thoi gian tao"),
            ("updated_at", "timestamp", "nullable", "Thoi gian cap nhat"),
        ],
    },
    {
        "name": "sport_types",
        "note": "Danh muc mon the thao de gan cho tung san.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma mon the thao"),
            ("name", "varchar(255)", "unique, not null", "Ten mon"),
            ("description", "text", "nullable", "Mo ta ngan"),
            ("created_at", "timestamp", "nullable", "Thoi gian tao"),
            ("updated_at", "timestamp", "nullable", "Thoi gian cap nhat"),
        ],
    },
    {
        "name": "courts",
        "note": "Thong tin san the thao va tinh trang khai thac.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma san"),
            ("sport_type_id", "bigint", "FK -> sport_types.id", "Mon the thao cua san"),
            ("name", "varchar(255)", "not null", "Ten san"),
            ("code", "varchar(255)", "unique, not null", "Ma san"),
            ("location", "varchar(255)", "not null", "Vi tri san"),
            ("capacity", "unsigned int", "not null", "Suc chua toi da"),
            ("status", "varchar(20)", "default active, index", "active/inactive/maintenance"),
            ("description", "text", "nullable", "Mo ta bo sung"),
        ],
    },
    {
        "name": "time_slots",
        "note": "Danh muc cac khung gio dat san.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma ca gio"),
            ("label", "varchar(255)", "unique, not null", "Nhan hien thi"),
            ("start_time", "time", "not null", "Gio bat dau"),
            ("end_time", "time", "not null", "Gio ket thuc"),
            ("created_at", "timestamp", "nullable", "Thoi gian tao"),
            ("updated_at", "timestamp", "nullable", "Thoi gian cap nhat"),
        ],
    },
    {
        "name": "court_schedules",
        "note": "Lich mo cua san theo thu trong tuan va ca gio.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma lich mo"),
            ("court_id", "bigint", "FK -> courts.id", "San duoc gan lich"),
            ("day_of_week", "tinyint", "not null", "Thu trong tuan"),
            ("time_slot_id", "bigint", "FK -> time_slots.id", "Ca gio duoc mo"),
            ("is_open", "boolean", "default true", "Co mo hay khong"),
            ("created_at", "timestamp", "nullable", "Thoi gian tao"),
            ("updated_at", "timestamp", "nullable", "Thoi gian cap nhat"),
        ],
    },
    {
        "name": "bookings",
        "note": "Bang header cua phieu dat, the hien nguoi dat va ket qua xu ly.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma booking"),
            ("user_id", "bigint", "FK -> users.id", "Nguoi tao booking"),
            ("purpose", "varchar(255)", "not null", "Muc dich dat san"),
            ("player_count", "unsigned int", "not null", "So nguoi tham gia"),
            ("contact_phone", "varchar(20)", "not null", "So dien thoai lien he"),
            ("status", "varchar(20)", "default pending, index", "pending/approved/rejected/cancelled/completed"),
            ("approved_by", "bigint", "FK -> users.id, nullable", "Admin duyet"),
            ("approved_at", "timestamp", "nullable", "Thoi diem duyet"),
            ("rejection_reason", "text", "nullable", "Ly do tu choi"),
            ("cancel_reason", "text", "nullable", "Ly do huy"),
            ("cancelled_at", "timestamp", "nullable", "Thoi diem huy"),
        ],
    },
    {
        "name": "booking_details",
        "note": "Chi tiet ngay dat, san dat va ca gio; dung de kiem tra trung lich.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma dong chi tiet"),
            ("booking_id", "bigint", "FK -> bookings.id", "Thuoc booking nao"),
            ("court_id", "bigint", "FK -> courts.id", "San duoc dat"),
            ("booking_date", "date", "index, not null", "Ngay su dung san"),
            ("time_slot_id", "bigint", "FK -> time_slots.id", "Ca gio dat"),
            ("created_at", "timestamp", "nullable", "Thoi gian tao"),
            ("updated_at", "timestamp", "nullable", "Thoi gian cap nhat"),
        ],
    },
    {
        "name": "usage_logs",
        "note": "Nhat ky ghi nhan ket qua su dung thuc te sau buoi dat san.",
        "columns": [
            ("id", "bigint", "PK, auto increment", "Ma nhat ky"),
            ("booking_detail_id", "bigint", "unique, FK -> booking_details.id", "Moi booking detail toi da mot nhat ky"),
            ("checked_by", "bigint", "FK -> users.id, nullable", "Nguoi cap nhat"),
            ("used_status", "varchar(20)", "default used", "used/no_show/cancelled"),
            ("checked_in_at", "timestamp", "nullable", "Moc bat dau"),
            ("checked_out_at", "timestamp", "nullable", "Moc ket thuc"),
            ("note", "text", "nullable", "Ghi chu bo sung"),
            ("created_at", "timestamp", "nullable", "Thoi gian tao"),
            ("updated_at", "timestamp", "nullable", "Thoi gian cap nhat"),
        ],
    },
]


def set_run_font(
    run,
    *,
    name: str = "Times New Roman",
    size: float = 13,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLOR_BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "bottom", "left", "right"):
        edge_el = tc_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def add_field(paragraph, instruction: str, *, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLOR_BLUE, 16, 8),
        ("Heading 2", 14, COLOR_NAVY, 10, 6),
        ("Heading 3", 13, COLOR_NAVY, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    prefix = paragraph.add_run("Bao cao de tai dat san the thao TLU - Trang ")
    set_run_font(prefix, size=10.5, color=COLOR_MUTED)
    add_field(paragraph, "PAGE", placeholder="1")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    size = {1: 16, 2: 14, 3: 13}[level]
    color = {1: COLOR_BLUE, 2: COLOR_NAVY, 3: COLOR_NAVY}[level]
    set_run_font(run, size=size, color=color, bold=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.35)
    run = paragraph.add_run(text)
    set_run_font(run, size=13, color=COLOR_NAVY)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        run = paragraph.add_run(item)
        set_run_font(run, size=13, color=COLOR_NAVY)


def add_numbered_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        run = paragraph.add_run(item)
        set_run_font(run, size=13, color=COLOR_NAVY)


def add_note_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [9200], table_width_dxa=9200, indent_dxa=120)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F9FC")
    set_cell_border(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, size=11.5, color=COLOR_NAVY)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple[str, ...]],
    width_weights: list[float],
    *,
    first_col_bold: bool = False,
    font_size: float = 10.8,
    header_fill: str = "E9F1F8",
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    widths = column_widths_from_weights(width_weights, 9200)
    apply_table_geometry(table, widths, table_width_dxa=9200, indent_dxa=120)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=11, color=COLOR_NAVY, bold=True)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(
                run,
                size=font_size,
                color=COLOR_NAVY,
                bold=(first_col_bold and col_idx == 0),
            )


def add_image(doc: Document, image_path: Path, caption: str, *, width: float = 6.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=11.5, color=COLOR_MUTED, italic=True)


def add_toc_page(doc: Document) -> None:
    add_heading(doc, "Muc luc", 1)
    add_note_box(
        doc,
        "Muc luc duoc sinh tu dong theo Heading. Khi mo file bang Microsoft Word, neu can co the bam Update Table de cap nhat so trang."
    )
    paragraph = doc.add_paragraph()
    add_field(paragraph, r'TOC \o "1-3" \h \z \u', placeholder="Cap nhat muc luc trong Word")
    doc.add_page_break()


def font_path(bold: bool = False) -> str:
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return ""


def load_font(size: int, *, bold: bool = False):
    path = font_path(bold)
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_text_to_width(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> str:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        test = " ".join(current + [word])
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] <= max_width:
            current.append(word)
        else:
            if current:
                lines.append(" ".join(current))
                current = [word]
            else:
                lines.append(word)
    if current:
        lines.append(" ".join(current))
    return "\n".join(lines)


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill, *, spacing: int = 6) -> None:
    wrapped = wrap_text_to_width(draw, text, font, box[2] - box[0] - 24)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=spacing, align="center")
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = box[0] + (box[2] - box[0] - text_width) / 2
    y = box[1] + (box[3] - box[1] - text_height) / 2
    draw.multiline_text((x, y), wrapped, font=font, fill=fill, spacing=spacing, align="center")


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    *,
    header_fill,
    body_fill,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=22, fill=body_fill, outline=COLOR_IMAGE_BLUE, width=3)
    header_box = (x1, y1, x2, y1 + 52)
    draw.rounded_rectangle(header_box, radius=22, fill=header_fill, outline=COLOR_IMAGE_BLUE, width=3)
    draw.rectangle((x1, y1 + 24, x2, y1 + 52), fill=header_fill, outline=header_fill)
    title_font = load_font(28, bold=True)
    body_font = load_font(23)
    draw_centered_text(draw, (x1 + 10, y1 + 6, x2 - 10, y1 + 48), title, title_font, COLOR_IMAGE_NAVY)
    draw_centered_text(draw, (x1 + 10, y1 + 60, x2 - 10, y2 - 10), body, body_font, COLOR_IMAGE_GRAY)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], *, color=COLOR_IMAGE_BLUE, width: int = 6) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head = 18
    left = (
        end[0] - head * math.cos(angle - math.pi / 6),
        end[1] - head * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - head * math.cos(angle + math.pi / 6),
        end[1] - head * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def draw_polyline_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], *, color=COLOR_IMAGE_BLUE, width: int = 5) -> None:
    for start, end in zip(points, points[1:]):
        draw.line([start, end], fill=color, width=width)
    draw_arrow(draw, points[-2], points[-1], color=color, width=width)


def create_use_case_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 920), COLOR_IMAGE_BG)
    draw = ImageDraw.Draw(image)
    title_font = load_font(38, bold=True)
    draw.text((460, 40), "So do tac nhan va nhom chuc nang chinh", font=title_font, fill=COLOR_IMAGE_NAVY)

    center_box = (500, 160, 1100, 760)
    draw_box(
        draw,
        center_box,
        "He thong dat san the thao TLU",
        "Quan ly tai nguyen\nQuan ly lich mo\nTiep nhan booking\nKiem tra xung dot\nPhe duyet va usage log\nThong ke su dung",
        header_fill=COLOR_IMAGE_LIGHT,
        body_fill=(255, 255, 255),
    )
    left_box = (80, 230, 420, 660)
    draw_box(
        draw,
        left_box,
        "Quan tri vien",
        "Dang nhap\nQuan ly mon the thao\nQuan ly san\nQuan ly ca gio va lich mo\nDuyet / tu choi booking\nCap nhat usage log\nXem bao cao",
        header_fill=COLOR_IMAGE_GREEN,
        body_fill=(255, 255, 255),
    )
    right_box = (1180, 260, 1520, 620)
    draw_box(
        draw,
        right_box,
        "Sinh vien",
        "Dang nhap\nTao booking\nXem booking cua minh\nSua / huy booking cua minh\nTheo doi trang thai xu ly",
        header_fill=COLOR_IMAGE_ORANGE,
        body_fill=(255, 255, 255),
    )
    draw_arrow(draw, (420, 445), (500, 445))
    draw_arrow(draw, (1180, 440), (1100, 440))
    note_font = load_font(22)
    note = "He thong dong thoi xu ly phan quyen, validation va kiem tra xung dot dat san."
    draw_centered_text(draw, (290, 790, 1310, 860), note, note_font, COLOR_IMAGE_GRAY)
    image.save(path)


def create_workflow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 760), COLOR_IMAGE_BG)
    draw = ImageDraw.Draw(image)
    title_font = load_font(38, bold=True)
    draw.text((580, 40), "Quy trinh nghiep vu tong the", font=title_font, fill=COLOR_IMAGE_NAVY)

    boxes = [
        ((60, 190, 280, 330), "1. Khai bao danh muc", "Mon the thao\nSan\nCa gio"),
        ((320, 190, 540, 330), "2. Lich mo", "Gan lich mo theo thu\nva tung ca gio"),
        ((580, 190, 800, 330), "3. Gui booking", "Sinh vien chon san,\nngay dat, ca gio"),
        ((840, 190, 1060, 330), "4. Kiem tra", "San hoat dong\nLich mo\nKhong trung lich"),
        ((1100, 190, 1320, 330), "5. Phe duyet", "Admin duyet\nhoac tu choi"),
        ((1360, 190, 1580, 330), "6. Usage log", "Danh dau used,\nno_show, cancelled"),
        ((710, 500, 1090, 650), "7. Bao cao va dashboard", "Tong booking, top san,\ntop sinh vien, xu huong su dung"),
    ]

    fills = [COLOR_IMAGE_LIGHT, COLOR_IMAGE_LIGHT, COLOR_IMAGE_ORANGE, COLOR_IMAGE_PINK, COLOR_IMAGE_GREEN, COLOR_IMAGE_GREEN, COLOR_IMAGE_LIGHT]
    for (box, title, body), fill in zip(boxes, fills):
        draw_box(draw, box, title, body, header_fill=fill, body_fill=(255, 255, 255))

    top_boxes = boxes[:6]
    for current, nxt in zip(top_boxes, top_boxes[1:]):
        draw_arrow(draw, (current[0][2], 260), (nxt[0][0], 260))
    draw_polyline_arrow(draw, [(1470, 330), (1470, 430), (900, 430), (900, 500)])
    image.save(path)


def create_mvc_diagram(path: Path) -> None:
    image = Image.new("RGB", (1700, 900), COLOR_IMAGE_BG)
    draw = ImageDraw.Draw(image)
    title_font = load_font(38, bold=True)
    draw.text((550, 40), "Kien truc MVC trong du an Laravel", font=title_font, fill=COLOR_IMAGE_NAVY)

    boxes = [
        ((80, 300, 300, 450), "Trinh duyet", "Nguoi dung thao tac voi giao dien web"),
        ((360, 300, 580, 450), "Routes", "Dinh tuyen request\nqua web.php"),
        ((640, 220, 930, 530), "Controllers", "Validation\nXu ly nghiep vu\nDieu huong va flash message"),
        ((1010, 220, 1280, 530), "Models", "Eloquent model\nRelationship\nTruy van du lieu"),
        ((1350, 220, 1620, 530), "MySQL", "Bookings\nCourts\nSchedules\nUsage logs"),
        ((640, 610, 930, 790), "Blade Views", "Dashboard\nForm CRUD\nBooking\nBao cao"),
    ]
    fills = [COLOR_IMAGE_ORANGE, COLOR_IMAGE_LIGHT, COLOR_IMAGE_LIGHT, COLOR_IMAGE_GREEN, COLOR_IMAGE_PINK, COLOR_IMAGE_LIGHT]
    for (box, title, body), fill in zip(boxes, fills):
        draw_box(draw, box, title, body, header_fill=fill, body_fill=(255, 255, 255))

    draw_arrow(draw, (300, 375), (360, 375))
    draw_arrow(draw, (580, 375), (640, 375))
    draw_arrow(draw, (930, 375), (1010, 375))
    draw_arrow(draw, (1280, 375), (1350, 375))
    draw_arrow(draw, (785, 530), (785, 610))
    draw_polyline_arrow(draw, [(785, 790), (785, 835), (190, 835), (190, 450)])

    note_font = load_font(22)
    draw_box(
        draw,
        (430, 120, 620, 210),
        "Middleware",
        "role:admin\nbao ve route quan tri",
        header_fill=COLOR_IMAGE_GREEN,
        body_fill=(255, 255, 255),
    )
    draw_box(
        draw,
        (970, 610, 1240, 790),
        "Policy va Auth",
        "BookingPolicy gioi han\nso huu booking\nDang nhap / dang xuat",
        header_fill=COLOR_IMAGE_ORANGE,
        body_fill=(255, 255, 255),
    )
    draw.text((1110, 842), "Luong request/response va cac lop bao ve nghiep vu", font=note_font, fill=COLOR_IMAGE_GRAY)
    image.save(path)


def create_module_diagram(path: Path) -> None:
    image = Image.new("RGB", (1750, 980), COLOR_IMAGE_BG)
    draw = ImageDraw.Draw(image)
    title_font = load_font(38, bold=True)
    draw.text((610, 40), "So do module he thong", font=title_font, fill=COLOR_IMAGE_NAVY)

    center_box = (690, 360, 1060, 610)
    draw_box(
        draw,
        center_box,
        "He thong dat san",
        "Xu ly booking\nKiem tra rang buoc\nDieu huong workflow",
        header_fill=COLOR_IMAGE_LIGHT,
        body_fill=(255, 255, 255),
    )

    modules = [
        ((120, 170, 470, 360), "Xac thuc", "Dang nhap\nDang xuat\nPhan quyen", COLOR_IMAGE_GREEN),
        ((120, 650, 470, 840), "Danh muc san", "Mon the thao\nSan\nCa gio", COLOR_IMAGE_LIGHT),
        ((560, 140, 910, 310), "Lich mo", "Gan lich mo theo thu\nva tung ca gio", COLOR_IMAGE_ORANGE),
        ((1180, 170, 1530, 360), "Booking", "Tao booking\nDuyet / tu choi\nHuy booking", COLOR_IMAGE_PINK),
        ((1180, 650, 1530, 840), "Usage log va bao cao", "Cap nhat su dung\nDashboard\nThong ke", COLOR_IMAGE_GREEN),
    ]
    for box, title, body, fill in modules:
        draw_box(draw, box, title, body, header_fill=fill, body_fill=(255, 255, 255))

    draw_arrow(draw, (470, 265), (690, 430))
    draw_arrow(draw, (470, 745), (690, 540))
    draw_arrow(draw, (860, 310), (860, 360))
    draw_arrow(draw, (1180, 265), (1060, 430))
    draw_arrow(draw, (1180, 745), (1060, 540))

    note_font = load_font(22)
    draw.text((450, 905), "Moi module tuong ung mot nhom chuc nang co lien ket du lieu va workflow ro rang.", font=note_font, fill=COLOR_IMAGE_GRAY)
    image.save(path)


def create_erd_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1240), COLOR_IMAGE_BG)
    draw = ImageDraw.Draw(image)
    title_font = load_font(38, bold=True)
    draw.text((560, 40), "So do ERD tong quan cua he thong", font=title_font, fill=COLOR_IMAGE_NAVY)

    boxes = {
        "users": ((70, 130, 420, 320), "users", "PK id\nname\nemail\nrole"),
        "sport_types": ((720, 130, 1070, 320), "sport_types", "PK id\nname\ndescription"),
        "time_slots": ((1380, 130, 1730, 320), "time_slots", "PK id\nlabel\nstart_time\nend_time"),
        "bookings": ((70, 430, 420, 680), "bookings", "PK id\nFK user_id\nstatus\napproved_by"),
        "courts": ((720, 430, 1070, 680), "courts", "PK id\nFK sport_type_id\ncode\nstatus"),
        "court_schedules": ((1380, 430, 1730, 680), "court_schedules", "PK id\nFK court_id\nFK time_slot_id\nday_of_week"),
        "booking_details": ((520, 840, 900, 1090), "booking_details", "PK id\nFK booking_id\nFK court_id\nFK time_slot_id\nbooking_date"),
        "usage_logs": ((1040, 840, 1420, 1090), "usage_logs", "PK id\nFK booking_detail_id\nFK checked_by\nused_status"),
    }

    fills = {
        "users": COLOR_IMAGE_GREEN,
        "sport_types": COLOR_IMAGE_LIGHT,
        "time_slots": COLOR_IMAGE_LIGHT,
        "bookings": COLOR_IMAGE_ORANGE,
        "courts": COLOR_IMAGE_LIGHT,
        "court_schedules": COLOR_IMAGE_PINK,
        "booking_details": COLOR_IMAGE_ORANGE,
        "usage_logs": COLOR_IMAGE_GREEN,
    }
    for key, (box, title, body) in boxes.items():
        draw_box(draw, box, title, body, header_fill=fills[key], body_fill=(255, 255, 255))

    draw_polyline_arrow(draw, [(245, 320), (245, 375), (245, 430)])
    draw_polyline_arrow(draw, [(895, 320), (895, 375), (895, 430)])
    draw_polyline_arrow(draw, [(1555, 320), (1555, 375), (1555, 430)])
    draw_polyline_arrow(draw, [(895, 680), (895, 760), (710, 760), (710, 840)])
    draw_polyline_arrow(draw, [(245, 680), (245, 760), (620, 760), (620, 840)])
    draw_polyline_arrow(draw, [(1555, 680), (1555, 760), (800, 760), (800, 840)])
    draw_polyline_arrow(draw, [(900, 965), (1040, 965)])
    draw_polyline_arrow(draw, [(245, 680), (245, 1140), (1230, 1140), (1230, 1090)])

    legend_font = load_font(22)
    legend = "Ngoai 8 bang nghiep vu cot loi, Laravel con co the tao them cac bang he thong nhu migrations, cache, jobs, sessions."
    draw_centered_text(draw, (240, 1145, 1560, 1210), legend, legend_font, COLOR_IMAGE_GRAY)
    image.save(path)


def create_ui_overview(path: Path) -> None:
    image = Image.new("RGB", (1850, 1080), (245, 248, 252))
    draw = ImageDraw.Draw(image)
    title_font = load_font(40, bold=True)
    draw.text((620, 35), "Minh hoa giao dien he thong", font=title_font, fill=COLOR_IMAGE_NAVY)

    panel_font = load_font(28, bold=True)
    small_font = load_font(20)
    tiny_font = load_font(17)

    panels = [
        ("Trang dang nhap", (70, 130, 570, 970)),
        ("Dashboard admin", (675, 130, 1175, 970)),
        ("Form dat san", (1280, 130, 1780, 970)),
    ]

    for label, box in panels:
        draw.rounded_rectangle(box, radius=28, fill=(255, 255, 255), outline=(196, 208, 224), width=3)
        draw.rounded_rectangle((box[0], box[1], box[2], box[1] + 64), radius=28, fill=(232, 240, 248), outline=(196, 208, 224), width=3)
        draw.rectangle((box[0], box[1] + 30, box[2], box[1] + 64), fill=(232, 240, 248), outline=(232, 240, 248))
        draw.text((box[0] + 24, box[1] + 16), label, font=panel_font, fill=COLOR_IMAGE_NAVY)

    # Login panel
    box = panels[0][1]
    draw.rounded_rectangle((box[0] + 70, box[1] + 150, box[2] - 70, box[1] + 720), radius=24, fill=(249, 251, 254), outline=(210, 220, 232), width=2)
    draw.text((box[0] + 165, box[1] + 200), "DANG NHAP HE THONG", font=panel_font, fill=COLOR_IMAGE_BLUE)
    for offset, label in [(320, "Email"), (450, "Mat khau")]:
        draw.text((box[0] + 110, box[1] + offset), label, font=small_font, fill=COLOR_IMAGE_GRAY)
        draw.rounded_rectangle((box[0] + 110, box[1] + offset + 40, box[2] - 110, box[1] + offset + 95), radius=16, fill=(255, 255, 255), outline=(192, 204, 220), width=2)
    draw.rounded_rectangle((box[0] + 110, box[1] + 585, box[2] - 110, box[1] + 650), radius=18, fill=(48, 108, 176), outline=(48, 108, 176))
    draw.text((box[0] + 195, box[1] + 602), "Dang nhap", font=panel_font, fill=(255, 255, 255))

    # Dashboard panel
    box = panels[1][1]
    draw.rounded_rectangle((box[0] + 30, box[1] + 95, box[0] + 140, box[3] - 30), radius=16, fill=(24, 53, 85), outline=(24, 53, 85))
    draw.text((box[0] + 48, box[1] + 130), "MENU", font=small_font, fill=(255, 255, 255))
    for idx, item in enumerate(["Tong quan", "Mon the thao", "San", "Lich mo", "Booking", "Bao cao"]):
        draw.text((box[0] + 48, box[1] + 200 + idx * 85), item, font=tiny_font, fill=(230, 238, 246))
    card_positions = [
        (box[0] + 175, box[1] + 120, box[0] + 325, box[1] + 245),
        (box[0] + 345, box[1] + 120, box[0] + 495, box[1] + 245),
        (box[0] + 175, box[1] + 270, box[0] + 325, box[1] + 395),
        (box[0] + 345, box[1] + 270, box[0] + 495, box[1] + 395),
    ]
    card_titles = ["Tong san", "Cho duyet", "Da duyet", "Su dung"]
    card_values = ["4", "1", "1", "1"]
    for pos, title, value in zip(card_positions, card_titles, card_values):
        draw.rounded_rectangle(pos, radius=18, fill=(246, 250, 254), outline=(198, 209, 223), width=2)
        draw.text((pos[0] + 18, pos[1] + 18), title, font=tiny_font, fill=COLOR_IMAGE_GRAY)
        draw.text((pos[0] + 55, pos[1] + 62), value, font=load_font(36, bold=True), fill=COLOR_IMAGE_BLUE)
    table_box = (box[0] + 175, box[1] + 440, box[2] - 30, box[3] - 55)
    draw.rounded_rectangle(table_box, radius=18, fill=(255, 255, 255), outline=(198, 209, 223), width=2)
    draw.rectangle((table_box[0], table_box[1], table_box[2], table_box[1] + 46), fill=(233, 241, 248), outline=(233, 241, 248))
    draw.text((table_box[0] + 18, table_box[1] + 12), "Danh sach booking gan day", font=tiny_font, fill=COLOR_IMAGE_NAVY)
    for y in [table_box[1] + 90, table_box[1] + 150, table_box[1] + 210]:
        draw.line((table_box[0] + 15, y, table_box[2] - 15, y), fill=(218, 225, 234), width=2)
    draw.text((table_box[0] + 22, table_box[1] + 62), "SV dat san", font=tiny_font, fill=COLOR_IMAGE_GRAY)
    draw.text((table_box[0] + 190, table_box[1] + 62), "San A1", font=tiny_font, fill=COLOR_IMAGE_GRAY)
    draw.text((table_box[0] + 320, table_box[1] + 62), "Cho duyet", font=tiny_font, fill=COLOR_IMAGE_GRAY)

    # Booking form panel
    box = panels[2][1]
    field_labels = [
        ("Muc dich su dung", 130),
        ("So nguoi tham gia", 235),
        ("So dien thoai", 340),
        ("Ngay dat", 445),
        ("Chon san", 550),
        ("Ca gio", 655),
    ]
    for label, offset in field_labels:
        draw.text((box[0] + 42, box[1] + offset), label, font=small_font, fill=COLOR_IMAGE_GRAY)
        draw.rounded_rectangle((box[0] + 42, box[1] + offset + 34, box[2] - 42, box[1] + offset + 88), radius=14, fill=(255, 255, 255), outline=(192, 204, 220), width=2)
    draw.rounded_rectangle((box[0] + 42, box[3] - 165, box[2] - 42, box[3] - 95), radius=18, fill=(48, 108, 176), outline=(48, 108, 176))
    draw.text((box[0] + 150, box[3] - 145), "Gui yeu cau dat san", font=panel_font, fill=(255, 255, 255))

    legend_font = load_font(22)
    draw.text((610, 1010), "Hinh minh hoa duoc dung de mo ta bo cuc cac man hinh chinh cua he thong.", font=legend_font, fill=COLOR_IMAGE_GRAY)
    image.save(path)


def generate_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "use_case": ASSET_DIR / "use-case.png",
        "workflow": ASSET_DIR / "workflow.png",
        "mvc": ASSET_DIR / "mvc.png",
        "module": ASSET_DIR / "module-diagram.png",
        "erd": ASSET_DIR / "erd-overview.png",
        "ui": ASSET_DIR / "ui-overview.png",
    }
    create_use_case_diagram(assets["use_case"])
    create_workflow_diagram(assets["workflow"])
    create_mvc_diagram(assets["mvc"])
    create_module_diagram(assets["module"])
    create_erd_diagram(assets["erd"])
    create_ui_overview(assets["ui"])
    return assets


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    add_footer(section)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(REPORT_META["institution"])
    set_run_font(run, size=16, color=COLOR_NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(REPORT_META["faculty"])
    set_run_font(run, size=14, color=COLOR_MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(REPORT_META["subject"].upper())
    set_run_font(run, size=19, color=COLOR_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(REPORT_META["title"].upper())
    set_run_font(run, size=20, color=COLOR_NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_META["topic_group"])
    set_run_font(run, size=13, color=COLOR_MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(REPORT_META["phase"])
    set_run_font(run, size=13, color=COLOR_MUTED, italic=True)

    rows = [
        ("De tai", REPORT_META["title"]),
        ("Cong nghe", REPORT_META["technology"]),
        ("Mo hinh kien truc", REPORT_META["architecture"]),
        ("Sinh vien thuc hien", REPORT_META["student_name"]),
        ("Ma sinh vien", REPORT_META["student_id"]),
        ("Lop", REPORT_META["class_name"]),
        ("Giang vien huong dan", REPORT_META["lecturer"]),
        ("Thoi gian", REPORT_META["date_label"]),
    ]
    add_table(doc, ["Thong tin", "Noi dung"], rows, [2.3, 5.7], first_col_bold=True, font_size=11.3)

    add_note_box(
        doc,
        "Luu y: truoc khi nop ban chinh thuc, sinh vien chi can dien thong tin ca nhan tren bia va kiem tra lai muc luc, so trang."
    )
    doc.add_page_break()


def add_preface(doc: Document) -> None:
    add_heading(doc, "Loi mo dau", 1)
    for paragraph in INTRO_PARAGRAPHS:
        add_body_paragraph(doc, paragraph)
    doc.add_page_break()


def add_chapter_one(doc: Document) -> None:
    add_heading(doc, "Chuong 1. Tong quan de tai", 1)
    add_heading(doc, "1.1. Ly do chon de tai", 2)
    add_body_paragraph(
        doc,
        "Trong khuon vien truong, nhu cau dat san bong da, bong chuyen, bong ro hoac cau long dien ra thuong xuyen. "
        "Neu viec quan ly van duoc thuc hien bang so tay, bang tinh hoac thong bao thu cong thi rat de xay ra trung lich, "
        "kho kiem soat tinh trang san va kho tong hop bao cao muc do su dung. Vi vay, mot he thong web tap trung la can thiet de chuan hoa quy trinh dat san."
    )

    add_heading(doc, "1.2. Muc tieu de tai", 2)
    add_bullets(doc, GOALS)

    add_heading(doc, "1.3. Doi tuong va pham vi nghien cuu", 2)
    add_body_paragraph(
        doc,
        "Doi tuong nghien cuu cua de tai gom he thong quan ly san the thao, quy trinh dat san cua sinh vien, "
        "nghiep vu phe duyet cua quan tri vien va du lieu thong ke su dung san trong khuon vien truong."
    )
    add_bullets(doc, SCOPES)

    add_heading(doc, "1.4. Phuong phap thuc hien", 2)
    add_bullets(doc, METHODS)


def add_chapter_two(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "Chuong 2. Phan tich yeu cau he thong", 1)
    add_heading(doc, "2.1. Boi canh nghiep vu", 2)
    add_body_paragraph(
        doc,
        "Bai toan dat ra la xay dung mot he thong cho phep nha truong quan ly danh muc san the thao, lich mo theo tung thu va tung ca gio, "
        "dong thoi ho tro sinh vien gui yeu cau dat san trong khung thoi gian hop le. He thong phai giai quyet duoc bai toan trung lich, "
        "the hien ro tinh trang san dang khai thac hay bao tri va luu vet ket qua su dung thuc te sau moi buoi dat san."
    )

    add_heading(doc, "2.2. Tac nhan su dung he thong", 2)
    add_table(doc, ["Tac nhan", "Quyen chinh", "Muc tieu su dung"], ROLE_ROWS, [1.5, 3.9, 2.6], first_col_bold=True, font_size=10.8)

    add_heading(doc, "2.3. Cac chuc nang chinh", 2)
    add_table(doc, ["Chuc nang", "Mo ta"], MAIN_FUNCTION_ROWS, [2.2, 5.8], first_col_bold=True, font_size=10.7)

    add_heading(doc, "2.4. Use Case tong quat", 2)
    add_image(doc, assets["use_case"], "Hinh 2.1. Use Case tong quat cua he thong", width=6.6)
    add_table(doc, ["Ma", "Use case", "Tac nhan", "Mo ta ngan"], USE_CASES, [0.9, 2.3, 1.5, 3.3], first_col_bold=True, font_size=10.3)

    add_heading(doc, "2.5. Workflow dat san", 2)
    add_body_paragraph(
        doc,
        "Workflow dat san la luong nghiep vu trung tam cua he thong, bat dau tu khai bao tai nguyen, tiep nhan yeu cau dat san, "
        "kiem tra rang buoc, phe duyet booking va ket thuc o usage log va bao cao."
    )
    add_image(doc, assets["workflow"], "Hinh 2.2. Workflow dat san tu khai bao tai nguyen den usage log", width=6.8)
    add_numbered_steps(doc, WORKFLOW_STEPS)

    add_heading(doc, "2.6. Yeu cau chuc nang", 2)
    add_bullets(doc, FUNCTIONAL_REQUIREMENTS)
    add_note_box(
        doc,
        "Rang buoc chuc nang quan trong: san bao tri khong duoc dat; booking khong duoc trung san - trung ngay - trung ca; "
        "sinh vien chi duoc thao tac tren booking cua chinh minh."
    )

    add_heading(doc, "2.7. Yeu cau phi chuc nang", 2)
    add_bullets(doc, NON_FUNCTIONAL_REQUIREMENTS)


def add_chapter_three(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "Chuong 3. Thiet ke he thong", 1)
    add_heading(doc, "3.1. So do kien truc MVC", 2)
    add_body_paragraph(
        doc,
        "Du an duoc xay dung tren Laravel theo mo hinh MVC. Request tu trinh duyet di qua route, middleware va controller; "
        "controller thuc hien validation, thao tac model va tra ket qua ve Blade View. Cach to chuc nay giup ma nguon de doc, de bao tri va phu hop voi yeu cau mon hoc."
    )
    add_image(doc, assets["mvc"], "Hinh 3.1. Kien truc MVC cua du an Laravel", width=6.7)

    add_heading(doc, "3.2. So do module", 2)
    add_body_paragraph(
        doc,
        "He thong duoc tach thanh cac module theo nhom nghiep vu chinh, giup viec phan chia chuc nang, ma nguon va du lieu tro nen ro rang hon."
    )
    add_image(doc, assets["module"], "Hinh 3.2. So do module cua he thong", width=6.8)
    add_table(doc, ["Module", "Noi dung"], MODULE_ROWS, [2.1, 5.9], first_col_bold=True, font_size=10.6)

    add_heading(doc, "3.3. Bieu do ERD", 2)
    add_image(doc, assets["erd"], "Hinh 3.3. Bieu do ERD tong quan cua he thong", width=6.9)

    add_heading(doc, "3.4. Thiet ke CSDL", 2)
    add_table(doc, ["Nhom bang", "Cac bang", "Y nghia"], TABLE_GROUPS, [1.8, 2.8, 3.4], first_col_bold=True, font_size=10.8)
    add_note_box(
        doc,
        "Ngoai 8 bang nghiep vu cot loi, trong phpMyAdmin co the xuat hien them cac bang nhu migrations, cache, jobs, sessions... Day la cac bang ha tang cua Laravel, khong phai bang nghiep vu chinh cua de tai."
    )
    add_table(doc, ["Bang nguon", "Bang dich", "Kieu", "Y nghia"], RELATIONSHIPS, [1.5, 2.3, 0.9, 3.1], first_col_bold=True, font_size=10.2)

    add_heading(doc, "3.5. Data dictionary", 2)
    add_body_paragraph(
        doc,
        "Phan nay tong hop cac cot quan trong, kieu du lieu va rang buoc chinh cua tung bang nghiep vu. "
        "Noi dung nay giup minh chung ro tieu chi A1, A2 va A4 trong barem cham diem."
    )
    for index, table_info in enumerate(DATA_DICTIONARY, start=1):
        add_heading(doc, f"3.5.{index}. Bang {table_info['name']}", 3)
        add_body_paragraph(doc, table_info["note"])
        add_table(
            doc,
            ["Cot", "Kieu du lieu", "Rang buoc", "Y nghia"],
            table_info["columns"],
            [1.7, 1.8, 2.2, 2.7],
            first_col_bold=True,
            font_size=10.0,
        )

    add_heading(doc, "3.6. Ma tran phan quyen", 2)
    add_table(doc, ["Chuc nang", "Admin", "Sinh vien", "Ghi chu"], PERMISSION_ROWS, [3.0, 1.0, 1.3, 2.9], first_col_bold=True, font_size=10.0)


def add_chapter_four(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "Chuong 4. Xay dung va cai dat he thong", 1)
    add_heading(doc, "4.1. Cong nghe su dung", 2)
    add_table(doc, ["Cong nghe", "Vai tro", "Ly do lua chon"], TECHNOLOGIES, [1.7, 2.2, 4.1], first_col_bold=True, font_size=10.5)

    add_heading(doc, "4.2. Moi truong cai dat", 2)
    add_table(doc, ["Thanh phan", "Gia tri", "Vai tro"], ENVIRONMENT_ROWS, [1.8, 1.8, 4.4], first_col_bold=True, font_size=10.3)

    add_heading(doc, "4.3. Cac chuc nang da hoan thanh", 2)
    add_table(doc, ["Nhom chuc nang", "Noi dung da xay dung"], COMPLETED_FEATURES, [2.4, 5.6], first_col_bold=True, font_size=10.7)
    add_note_box(
        doc,
        "Tai khoan demo hien co: admin@campus.local / password, student1@campus.local / password, student2@campus.local / password."
    )

    add_heading(doc, "4.4. Minh hoa giao dien", 2)
    add_body_paragraph(
        doc,
        "De phuc vu bao cao, duoi day la hinh minh hoa cho ba man hinh dai dien cua he thong gom trang dang nhap, dashboard quan tri va form dat san."
    )
    add_image(doc, assets["ui"], "Hinh 4.1. Minh hoa mot so man hinh chinh cua he thong", width=6.9)
    add_note_box(
        doc,
        "Khi demo truc tiep, co the dang nhap bang tai khoan admin de minh hoa dashboard, CRUD danh muc va phe duyet booking; sau do dang nhap bang tai khoan sinh vien de minh hoa form dat san."
    )


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "Ket luan", 1)
    add_body_paragraph(
        doc,
        "De tai \"Xay dung he thong quan ly va dat san the thao tai TLU\" da dat duoc muc tieu chinh trong pham vi vong 1: "
        "hinh thanh mot he thong web co quy trinh ro rang, co kiem soat xung dot, co phan quyen nguoi dung va co kha nang bao cao su dung san. "
        "Ket qua nay tao nen mot nen tang tot de tiep tuc mo rong them cac tinh nang nang cao trong cac vong sau, dong thoi dap ung phan lon tieu chi diem nen cua mon hoc."
    )
    add_heading(doc, "Danh gia tong quat", 2)
    add_note_box(
        doc,
        "Muc do hoan thanh hien tai co the xem la da dap ung tot nhom tieu chi A1 den A8. Hai noi dung can hoan thien them truoc khi nop la "
        "bo sung thong tin ca nhan tren bia, chot commit minh chung theo moc va ra soat README / ERD de dong bo voi ban nop cuoi."
    )
    add_heading(doc, "Han che hien tai", 2)
    add_bullets(doc, LIMITATIONS)
    add_heading(doc, "Huong phat trien", 2)
    add_bullets(doc, FUTURE_DIRECTIONS)


def build_report() -> Document:
    assets = generate_assets()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_preface(doc)
    add_toc_page(doc)
    add_chapter_one(doc)
    add_chapter_two(doc, assets)
    add_chapter_three(doc, assets)
    add_chapter_four(doc, assets)
    add_conclusion(doc)
    return doc


def main() -> None:
    doc = build_report()
    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")
    try:
        doc.save(LEGACY_OUTPUT_PATH)
        print(f"Synced: {LEGACY_OUTPUT_PATH}")
    except PermissionError:
        print(f"Warning: could not overwrite locked file: {LEGACY_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
