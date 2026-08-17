from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from table_geometry import apply_table_geometry, column_widths_from_weights


BASE_DIR = Path(__file__).resolve().parents[1]
ASSET_DIR = BASE_DIR / "generated" / "report-assets"
OUTPUT_PATH = BASE_DIR / "bao-cao-de-tai-dat-san-the-thao-co-dau.docx"

COLOR_NAVY = RGBColor(0x12, 0x2A, 0x43)
COLOR_BLUE = RGBColor(0x1F, 0x5A, 0x94)
COLOR_MUTED = RGBColor(0x5B, 0x66, 0x73)
COLOR_BORDER = "C9D4E1"
COLOR_HEADER = "E9F1F8"
COLOR_NOTE = "F5F8FC"

META_ROWS = [
    ("Đề tài", "Xây dựng hệ thống quản lý và đặt sân thể thao tại TLU"),
    ("Công nghệ", "Laravel 12, PHP 8.2, Blade, Bootstrap 5, MySQL, phpMyAdmin"),
    ("Mô hình kiến trúc", "MVC"),
    ("Sinh viên thực hiện", "[Điền họ tên]"),
    ("Mã sinh viên", "[Điền mã sinh viên]"),
    ("Lớp", "[Điền lớp]"),
    ("Giảng viên hướng dẫn", "[Điền tên giảng viên hướng dẫn]"),
    ("Thời gian", "Tháng 08 năm 2026"),
]

INTRO_PARAGRAPHS = [
    "Trong bối cảnh chuyển đổi số trong giáo dục, nhu cầu tin học hóa các dịch vụ hỗ trợ sinh viên ngày càng trở nên cần thiết. Bên cạnh các nghiệp vụ học tập, nhóm dịch vụ trong khuôn viên như quản lý và đặt sân thể thao cũng cần được tổ chức bài bản để hạn chế tình trạng trùng lịch, khó theo dõi tình trạng sân và khó tổng hợp báo cáo sử dụng.",
    "Từ thực tế đó, đề tài “Xây dựng hệ thống quản lý và đặt sân thể thao tại TLU” được lựa chọn nhằm xây dựng một hệ thống web hỗ trợ quản lý danh mục sân, lịch mở, tiếp nhận yêu cầu đặt sân, kiểm tra xung đột, phân quyền người dùng và thống kê quá trình sử dụng sân thể thao trong khuôn viên trường.",
    "Báo cáo này trình bày quá trình phân tích yêu cầu, thiết kế cơ sở dữ liệu, xây dựng hệ thống theo mô hình MVC và đánh giá kết quả đã đạt được trong phạm vi vòng 1. Nội dung được sắp xếp đúng theo khung báo cáo môn học để có thể dùng trực tiếp cho hồ sơ nộp bài.",
]

GOALS = [
    "Số hóa quy trình quản lý và đặt sân thể thao trong khuôn viên trường.",
    "Hạn chế tình trạng trùng lịch giữa các yêu cầu đặt sân cùng ngày và cùng ca giờ.",
    "Cho phép theo dõi tình trạng sân, lịch mở, booking, usage log và báo cáo sử dụng.",
    "Phân quyền rõ ràng giữa admin và sinh viên theo đúng workflow nghiệp vụ.",
    "Đảm bảo dự án có thể cài đặt lại nhanh bằng migration, seeder và dữ liệu demo.",
]

SCOPES = [
    "Đối tượng nghiên cứu là quy trình quản lý sân thể thao, lịch mở, quy trình đặt sân của sinh viên, phê duyệt booking của quản trị viên và thống kê sử dụng sân.",
    "Phiên bản hiện tại triển khai 8 bảng nghiệp vụ cốt lõi: users, sport_types, courts, time_slots, court_schedules, bookings, booking_details, usage_logs.",
    "Hệ thống chưa triển khai thanh toán, thông báo email và quản lý lịch bảo trì chi tiết theo từng khoảng thời gian.",
    "Phạm vi thực hiện là project cá nhân chạy local, quản lý mã nguồn bằng Git và đẩy lên GitHub.",
]

METHODS = [
    "Khảo sát bài toán và xác định quy trình nghiệp vụ đặt sân theo vai trò người dùng.",
    "Phân tích tác nhân, use case, quy tắc ràng buộc và các tình huống xung đột cần xử lý.",
    "Thiết kế cơ sở dữ liệu theo quan hệ cha - con và mô hình MVC cho ứng dụng Laravel.",
    "Xây dựng migration, seeder, model, controller, blade view và phần xác thực, phân quyền.",
    "Kiểm thử bằng dữ liệu demo, kiểm tra đăng nhập, booking, duyệt, hủy và thống kê.",
]

ROLE_ROWS = [
    ("Admin", "Đăng nhập, CRUD danh mục, duyệt booking, cập nhật usage log, xem báo cáo", "Quản lý toàn bộ tài nguyên và workflow hệ thống"),
    ("Sinh viên", "Đăng nhập, tạo booking, xem, sửa, hủy booking của chính mình", "Sử dụng sân thể thao phù hợp nhu cầu học tập và sinh hoạt"),
]

MAIN_FUNCTION_ROWS = [
    ("Đăng nhập và phân quyền", "Xác thực người dùng và hiển thị đúng phạm vi chức năng theo vai trò admin / sinh viên."),
    ("Quản lý danh mục sân", "Quản lý môn thể thao, sân, ca giờ và tình trạng sân."),
    ("Quản lý lịch mở", "Gán lịch mở cho từng sân theo thứ trong tuần và từng ca giờ."),
    ("Đặt sân", "Sinh viên tạo phiếu đặt, nhập mục đích, số người, ngày đặt và ca giờ."),
    ("Phê duyệt và hủy booking", "Admin duyệt / từ chối booking; sinh viên được hủy booking của mình khi còn hợp lệ."),
    ("Usage log và báo cáo", "Cập nhật kết quả sử dụng thực tế và tổng hợp dashboard, thống kê."),
]

USE_CASE_ROWS = [
    ("UC01", "Đăng nhập hệ thống", "Admin, Sinh viên", "Người dùng xác thực để truy cập hệ thống theo vai trò."),
    ("UC02", "Quản lý môn thể thao", "Admin", "Thêm, sửa, xóa danh mục môn thể thao phục vụ khai báo sân."),
    ("UC03", "Quản lý sân", "Admin", "Thêm, sửa, xóa sân; cập nhật sức chứa, vị trí và trạng thái."),
    ("UC04", "Quản lý ca giờ", "Admin", "Khai báo các khung giờ sử dụng cho đặt sân."),
    ("UC05", "Quản lý lịch mở sân", "Admin", "Gán lịch mở theo thứ trong tuần và từng ca giờ cho mỗi sân."),
    ("UC06", "Tạo booking", "Sinh viên", "Nhập thông tin đặt sân và gửi yêu cầu đến hệ thống."),
    ("UC07", "Kiểm tra xung đột", "Hệ thống", "Chặn đặt sân nếu trùng lịch, sân bảo trì hoặc không mở lịch."),
    ("UC08", "Duyệt hoặc từ chối booking", "Admin", "Cập nhật trạng thái booking và lý do từ chối nếu có."),
    ("UC09", "Hủy booking", "Sinh viên", "Sinh viên hủy phiếu đặt của chính mình khi chưa hoàn tất."),
    ("UC10", "Cập nhật usage log", "Admin", "Ghi nhận used, no_show hoặc cancelled sau buổi đặt sân."),
    ("UC11", "Xem báo cáo", "Admin", "Tổng hợp thống kê booking, usage log, top sân và top sinh viên."),
]

WORKFLOW_STEPS = [
    "Admin khai báo môn thể thao, sân, ca giờ và lịch mở.",
    "Sinh viên chọn sân, ngày đặt, ca giờ và nhập thông tin booking.",
    "Hệ thống kiểm tra tình trạng sân, sức chứa, lịch mở và xung đột.",
    "Nếu hợp lệ, booking được tạo ở trạng thái pending.",
    "Admin duyệt hoặc từ chối booking.",
    "Sau buổi đặt sân, admin cập nhật usage log.",
    "Dashboard và báo cáo tổng hợp dữ liệu sử dụng sân.",
]

FUNCTIONAL_REQUIREMENTS = [
    "Đăng nhập và đăng xuất hệ thống theo tài khoản đã cấp.",
    "Quản lý danh mục môn thể thao, sân, ca giờ và lịch mở của từng sân.",
    "Cho phép sinh viên tạo phiếu đặt sân với thông tin mục đích, số người, ngày đặt và ca giờ.",
    "Kiểm tra xung đột booking theo sân, ngày, ca giờ và tình trạng hoạt động của sân.",
    "Cho phép admin duyệt hoặc từ chối booking, lưu lý do từ chối khi cần.",
    "Cho phép sinh viên xem, sửa, hủy booking của chính mình khi còn hợp lệ.",
    "Cập nhật usage log sau buổi đặt sân để đồng bộ trạng thái booking.",
    "Tổng hợp dashboard và báo cáo thống kê theo khoảng thời gian.",
]

NON_FUNCTIONAL_REQUIREMENTS = [
    "Giao diện dễ hiểu, viết bằng tiếng Việt và có bố cục rõ ràng cho form, bảng và dashboard.",
    "Dữ liệu được validate ở backend; không thao tác thay đổi dữ liệu bằng GET.",
    "Mật khẩu được mã hóa; route quản trị được bảo vệ bởi middleware và policy.",
    "Hệ thống có thể cài đặt lại bằng migrate:fresh --seed để phục vụ demo nhanh.",
    "Cấu trúc MVC rõ ràng, View không query database và Controller không chứa SQL dài.",
]

MODULE_ROWS = [
    ("Module xác thực", "Đăng nhập, đăng xuất và phân quyền theo vai trò."),
    ("Module danh mục", "Quản lý môn thể thao, sân, ca giờ và trạng thái sân."),
    ("Module lịch mở", "Khai báo lịch mở cho từng sân theo thứ và khung giờ."),
    ("Module booking", "Tạo, xem, sửa, hủy, duyệt và từ chối booking."),
    ("Module usage log", "Cập nhật used, no_show, cancelled cho booking detail."),
    ("Module báo cáo", "Thống kê booking, top sân, top sinh viên và xu hướng sử dụng."),
]

TABLE_GROUPS = [
    ("Nhóm tài khoản", "users", "Lưu thông tin đăng nhập, vai trò và liên kết người tạo booking."),
    ("Nhóm danh mục", "sport_types, courts, time_slots", "Quản lý môn thể thao, sân và khung giờ."),
    ("Nhóm lịch", "court_schedules", "Quy định sân nào mở vào thứ nào, ca giờ nào."),
    ("Nhóm giao dịch", "bookings, booking_details", "Lưu phiếu đặt và chi tiết ngày đặt, sân đặt, ca giờ."),
    ("Nhóm theo dõi", "usage_logs", "Ghi nhận kết quả sử dụng thực tế sau buổi đặt sân."),
]

RELATIONSHIP_ROWS = [
    ("users", "bookings", "1 - N", "Một người dùng có thể tạo nhiều phiếu đặt."),
    ("users", "bookings (approved_by)", "1 - N", "Một admin có thể duyệt nhiều booking."),
    ("sport_types", "courts", "1 - N", "Một môn thể thao có nhiều sân."),
    ("courts", "court_schedules", "1 - N", "Một sân có nhiều lịch mở theo thứ và ca."),
    ("time_slots", "court_schedules", "1 - N", "Một ca giờ có thể áp dụng cho nhiều sân."),
    ("bookings", "booking_details", "1 - N", "Một phiếu đặt có thể mở rộng nhiều dòng chi tiết."),
    ("courts", "booking_details", "1 - N", "Một sân xuất hiện trong nhiều booking detail."),
    ("time_slots", "booking_details", "1 - N", "Một ca giờ xuất hiện trong nhiều booking detail."),
    ("booking_details", "usage_logs", "1 - 1", "Mỗi dòng chi tiết có tối đa một usage log."),
    ("users", "usage_logs", "1 - N", "Admin cập nhật usage log cho từng buổi sử dụng."),
]

PERMISSION_ROWS = [
    ("Đăng nhập / đăng xuất", "Có", "Có", "Tất cả người dùng đã cấp tài khoản đều sử dụng được."),
    ("CRUD môn thể thao", "Có", "Không", "Chỉ admin quản lý danh mục sport_types."),
    ("CRUD sân", "Có", "Không", "Chỉ admin quản lý courts."),
    ("CRUD ca giờ và lịch mở", "Có", "Không", "Chỉ admin quản lý time_slots và court_schedules."),
    ("Tạo booking", "Không", "Có", "Sinh viên chủ động gửi yêu cầu đặt sân."),
    ("Xem tất cả booking", "Có", "Không", "Admin theo dõi toàn bộ workflow."),
    ("Xem booking của mình", "Không", "Có", "Policy giới hạn theo user_id."),
    ("Sửa / hủy booking của mình", "Không", "Có", "Chỉ khi booking chưa ở trạng thái cancelled/completed."),
    ("Duyệt / từ chối booking", "Có", "Không", "Admin xử lý workflow pending."),
    ("Cập nhật usage log", "Có", "Không", "Admin ghi nhận kết quả sử dụng."),
    ("Xem dashboard và báo cáo", "Có", "Có một phần", "Sinh viên xem dashboard cá nhân; báo cáo tổng hợp dành cho admin."),
]

TECHNOLOGY_ROWS = [
    ("PHP 8.2", "Ngôn ngữ lập trình backend", "Phù hợp Laravel và dễ triển khai local trên XAMPP."),
    ("Laravel 12", "Framework MVC", "Hỗ trợ route, middleware, Eloquent, validation và Blade."),
    ("Blade", "Template engine", "Xây dựng giao diện dashboard, form CRUD và trang booking."),
    ("Bootstrap 5", "CSS framework", "Giúp giao diện responsive và thống nhất bố cục."),
    ("MySQL", "Hệ quản trị CSDL", "Lưu dữ liệu booking, danh mục, lịch mở và usage log."),
    ("phpMyAdmin", "Quản lý CSDL", "Quan sát bảng và dữ liệu local trong quá trình demo."),
    ("Git và GitHub", "Quản lý mã nguồn", "Lưu vết các mốc hoàn thiện và đẩy dự án lên kho mã nguồn."),
]

ENVIRONMENT_ROWS = [
    ("Hệ điều hành", "Windows", "Môi trường phát triển local trên máy cá nhân."),
    ("PHP", "8.2.12", "Chạy Laravel thông qua XAMPP / CLI."),
    ("Composer", "2.10.2", "Quản lý dependency PHP."),
    ("Node.js", "v24.16.0", "Build tài nguyên frontend bằng Vite."),
    ("npm", "11.13.0", "Cài đặt package frontend."),
    ("Web server", "Apache (XAMPP)", "Phục vụ ứng dụng local."),
    ("Database", "MySQL (XAMPP)", "Lưu dữ liệu nghiệp vụ của hệ thống."),
    ("Công cụ CSDL", "phpMyAdmin", "Quan sát schema và dữ liệu demo."),
    ("Terminal", "Git Bash", "Thực hiện các lệnh cài đặt, git và artisan."),
]

COMPLETED_FEATURE_ROWS = [
    ("Xác thực và phân quyền", "Đăng nhập, đăng xuất, middleware role:admin và BookingPolicy."),
    ("Quản lý tài nguyên", "CRUD sport_types, courts, time_slots, court_schedules."),
    ("Workflow booking", "Tạo booking, kiểm tra xung đột, xem, sửa, hủy, duyệt và từ chối."),
    ("Usage log", "Ghi nhận used, no_show, cancelled và đồng bộ trạng thái booking."),
    ("Dashboard và báo cáo", "Thống kê tổng quan theo vai trò, top sân, top sinh viên và xu hướng đặt sân."),
    ("Dữ liệu demo", "Seeder tài khoản, sân, lịch mở, booking và usage log để demo nhanh."),
]

LIMITATIONS = [
    "Chưa có thông báo email hoặc hệ thống nhắc lịch sau khi booking được duyệt.",
    "Chưa triển khai giới hạn số lần đặt sân theo tuần cho từng sinh viên.",
    "Chưa tách lịch bảo trì thành một module riêng có khung thời gian cụ thể.",
    "Chưa xuất báo cáo sang PDF/Excel ngay trong giao diện web.",
]

FUTURE_DIRECTIONS = [
    "Bổ sung booking policy nâng cao: giới hạn số slot / sinh viên / tuần.",
    "Thêm module bảo trì sân theo khoảng thời gian để khóa sân tạm thời.",
    "Tích hợp email, toast notification hoặc thông báo trong hệ thống.",
    "Bổ sung xuất báo cáo PDF/Excel và lịch sử thao tác audit log.",
    "Mở rộng đặt nhiều ca giờ trong cùng một phiếu booking nếu đề tài yêu cầu.",
]

DATA_DICTIONARY = [
    ("users", "Bảng lưu thông tin tài khoản đăng nhập và vai trò người dùng.", [
        ("id", "bigint", "PK, auto increment", "Mã người dùng"),
        ("name", "varchar(255)", "not null", "Họ tên hiển thị"),
        ("email", "varchar(255)", "unique, not null", "Email đăng nhập"),
        ("password", "varchar(255)", "not null", "Mật khẩu đã hash"),
        ("role", "varchar(20)", "default student, index", "Vai trò admin/student"),
        ("remember_token", "varchar(100)", "nullable", "Token ghi nhớ đăng nhập"),
        ("created_at", "timestamp", "nullable", "Thời gian tạo"),
        ("updated_at", "timestamp", "nullable", "Thời gian cập nhật"),
    ]),
    ("sport_types", "Danh mục môn thể thao để gán cho từng sân.", [
        ("id", "bigint", "PK, auto increment", "Mã môn thể thao"),
        ("name", "varchar(255)", "unique, not null", "Tên môn"),
        ("description", "text", "nullable", "Mô tả ngắn"),
        ("created_at", "timestamp", "nullable", "Thời gian tạo"),
        ("updated_at", "timestamp", "nullable", "Thời gian cập nhật"),
    ]),
    ("courts", "Thông tin sân thể thao và tình trạng khai thác.", [
        ("id", "bigint", "PK, auto increment", "Mã sân"),
        ("sport_type_id", "bigint", "FK -> sport_types.id", "Môn thể thao của sân"),
        ("name", "varchar(255)", "not null", "Tên sân"),
        ("code", "varchar(255)", "unique, not null", "Mã sân"),
        ("location", "varchar(255)", "not null", "Vị trí sân"),
        ("capacity", "unsigned int", "not null", "Sức chứa tối đa"),
        ("status", "varchar(20)", "default active, index", "active/inactive/maintenance"),
        ("description", "text", "nullable", "Mô tả bổ sung"),
    ]),
    ("time_slots", "Danh mục các khung giờ đặt sân.", [
        ("id", "bigint", "PK, auto increment", "Mã ca giờ"),
        ("label", "varchar(255)", "unique, not null", "Nhãn hiển thị"),
        ("start_time", "time", "not null", "Giờ bắt đầu"),
        ("end_time", "time", "not null", "Giờ kết thúc"),
        ("created_at", "timestamp", "nullable", "Thời gian tạo"),
        ("updated_at", "timestamp", "nullable", "Thời gian cập nhật"),
    ]),
    ("court_schedules", "Lịch mở của sân theo thứ trong tuần và ca giờ.", [
        ("id", "bigint", "PK, auto increment", "Mã lịch mở"),
        ("court_id", "bigint", "FK -> courts.id", "Sân được gán lịch"),
        ("day_of_week", "tinyint", "not null", "Thứ trong tuần"),
        ("time_slot_id", "bigint", "FK -> time_slots.id", "Ca giờ được mở"),
        ("is_open", "boolean", "default true", "Có mở hay không"),
        ("created_at", "timestamp", "nullable", "Thời gian tạo"),
        ("updated_at", "timestamp", "nullable", "Thời gian cập nhật"),
    ]),
    ("bookings", "Bảng header của phiếu đặt, thể hiện người đặt và kết quả xử lý.", [
        ("id", "bigint", "PK, auto increment", "Mã booking"),
        ("user_id", "bigint", "FK -> users.id", "Người tạo booking"),
        ("purpose", "varchar(255)", "not null", "Mục đích đặt sân"),
        ("player_count", "unsigned int", "not null", "Số người tham gia"),
        ("contact_phone", "varchar(20)", "not null", "Số điện thoại liên hệ"),
        ("status", "varchar(20)", "default pending, index", "pending/approved/rejected/cancelled/completed"),
        ("approved_by", "bigint", "FK -> users.id, nullable", "Admin duyệt"),
        ("approved_at", "timestamp", "nullable", "Thời điểm duyệt"),
        ("rejection_reason", "text", "nullable", "Lý do từ chối"),
        ("cancel_reason", "text", "nullable", "Lý do hủy"),
        ("cancelled_at", "timestamp", "nullable", "Thời điểm hủy"),
    ]),
    ("booking_details", "Chi tiết ngày đặt, sân đặt và ca giờ; dùng để kiểm tra trùng lịch.", [
        ("id", "bigint", "PK, auto increment", "Mã dòng chi tiết"),
        ("booking_id", "bigint", "FK -> bookings.id", "Thuộc booking nào"),
        ("court_id", "bigint", "FK -> courts.id", "Sân được đặt"),
        ("booking_date", "date", "index, not null", "Ngày sử dụng sân"),
        ("time_slot_id", "bigint", "FK -> time_slots.id", "Ca giờ đặt"),
        ("created_at", "timestamp", "nullable", "Thời gian tạo"),
        ("updated_at", "timestamp", "nullable", "Thời gian cập nhật"),
    ]),
    ("usage_logs", "Nhật ký ghi nhận kết quả sử dụng thực tế sau buổi đặt sân.", [
        ("id", "bigint", "PK, auto increment", "Mã nhật ký"),
        ("booking_detail_id", "bigint", "unique, FK -> booking_details.id", "Mỗi booking detail tối đa một nhật ký"),
        ("checked_by", "bigint", "FK -> users.id, nullable", "Người cập nhật"),
        ("used_status", "varchar(20)", "default used", "used/no_show/cancelled"),
        ("checked_in_at", "timestamp", "nullable", "Mốc bắt đầu"),
        ("checked_out_at", "timestamp", "nullable", "Mốc kết thúc"),
        ("note", "text", "nullable", "Ghi chú bổ sung"),
        ("created_at", "timestamp", "nullable", "Thời gian tạo"),
        ("updated_at", "timestamp", "nullable", "Thời gian cập nhật"),
    ]),
]


def set_run_font(run, *, size: float = 13, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None, name: str = "Times New Roman") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_border(cell, color: str = COLOR_BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "bottom", "left", "right"):
        edge_el = borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, COLOR_BLUE),
        ("Heading 2", 14, COLOR_NAVY),
        ("Heading 3", 13, COLOR_NAVY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = placeholder
    separate.append(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_footer(doc: Document) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Báo cáo đề tài đặt sân thể thao TLU - Trang ")
    set_run_font(run, size=10.5, color=COLOR_MUTED)
    add_field(p, "PAGE", "1")


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 14, 3: 13}[level], bold=True, color={1: COLOR_BLUE, 2: COLOR_NAVY, 3: COLOR_NAVY}[level])


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    run = p.add_run(text)
    set_run_font(run, size=13, color=COLOR_NAVY)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        set_run_font(run, size=13, color=COLOR_NAVY)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        set_run_font(run, size=13, color=COLOR_NAVY)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], weights: list[float], *, first_col_bold: bool = False, font_size: float = 10.8) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    widths = column_widths_from_weights(weights, 9200)
    apply_table_geometry(table, widths, table_width_dxa=9200, indent_dxa=120)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, COLOR_HEADER)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=11, bold=True, color=COLOR_NAVY)

    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=COLOR_NAVY, bold=(first_col_bold and c == 0))


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [9200], table_width_dxa=9200, indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLOR_NOTE)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=11.5, color=COLOR_NAVY)


def add_image(doc: Document, key: str, caption: str, *, width: float = 6.8) -> None:
    image_path = ASSET_DIR / key
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(width))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        set_run_font(run, size=11.5, italic=True, color=COLOR_MUTED)
    else:
        add_note(doc, f"Chưa tìm thấy hình minh họa: {image_path.name}")


def add_toc(doc: Document) -> None:
    add_heading(doc, "Mục lục", 1)
    add_note(doc, "Mục lục được sinh tự động theo Heading. Nếu mở bằng Microsoft Word, có thể bấm Update Table để cập nhật số trang.")
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u', "Cập nhật mục lục trong Word")
    doc.add_page_break()


def build_report() -> Document:
    doc = Document()
    style_document(doc)
    add_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC THỦY LỢI (TLU)")
    set_run_font(run, size=16, bold=True, color=COLOR_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_run_font(run, size=14, bold=True, color=COLOR_MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("BÁO CÁO PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG HỆ THỐNG")
    set_run_font(run, size=19, bold=True, color=COLOR_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("XÂY DỰNG HỆ THỐNG QUẢN LÝ VÀ ĐẶT SÂN THỂ THAO TẠI TLU")
    set_run_font(run, size=20, bold=True, color=COLOR_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Đề tài 6 - Nhóm 2 - Dịch vụ sinh viên và khuôn viên")
    set_run_font(run, size=13, italic=True, color=COLOR_MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phạm vi triển khai: Vòng 1")
    set_run_font(run, size=13, italic=True, color=COLOR_MUTED)

    add_table(doc, ["Thông tin", "Nội dung"], META_ROWS, [2.3, 5.7], first_col_bold=True, font_size=11.3)
    add_note(doc, "Lưu ý: trước khi nộp bản chính thức, sinh viên cần điền đầy đủ thông tin cá nhân trên bìa và kiểm tra lại mục lục.")
    doc.add_page_break()

    add_heading(doc, "Lời mở đầu", 1)
    for paragraph in INTRO_PARAGRAPHS:
        add_body(doc, paragraph)
    doc.add_page_break()

    add_toc(doc)

    add_heading(doc, "CHƯƠNG I: Tổng Quan Đề Tài", 1)
    add_heading(doc, "1. Lý do chọn đề tài", 2)
    add_body(doc, "Trong khuôn viên trường, nhu cầu đặt sân bóng đá, bóng chuyền, bóng rổ hoặc cầu lông diễn ra thường xuyên. Nếu việc quản lý vẫn được thực hiện bằng sổ tay, bảng tính hoặc thông báo thủ công thì rất dễ xảy ra trùng lịch, khó kiểm soát tình trạng sân và khó tổng hợp báo cáo mức độ sử dụng. Vì vậy, một hệ thống web tập trung là cần thiết để chuẩn hóa quy trình đặt sân.")
    add_heading(doc, "2. Mục tiêu đề tài", 2)
    add_bullets(doc, GOALS)
    add_heading(doc, "3. Đối tượng và phạm vi nghiên cứu", 2)
    add_bullets(doc, SCOPES)
    add_heading(doc, "4. Phương pháp thực hiện", 2)
    add_bullets(doc, METHODS)

    add_heading(doc, "CHƯƠNG II: Phân tích yêu cầu hệ thống", 1)
    add_heading(doc, "1. Bối cảnh nghiệp vụ", 2)
    add_body(doc, "Bài toán đặt ra là xây dựng một hệ thống cho phép nhà trường quản lý danh mục sân thể thao, lịch mở theo từng thứ và từng ca giờ, đồng thời hỗ trợ sinh viên gửi yêu cầu đặt sân trong khung thời gian hợp lệ. Hệ thống phải giải quyết được bài toán trùng lịch, thể hiện rõ tình trạng sân đang khai thác hay bảo trì và lưu vết kết quả sử dụng thực tế sau mỗi buổi đặt sân.")
    add_heading(doc, "2. Tác nhân sử dụng hệ thống", 2)
    add_table(doc, ["Tác nhân", "Quyền chính", "Mục tiêu sử dụng"], ROLE_ROWS, [1.5, 3.9, 2.6], first_col_bold=True)
    add_heading(doc, "3. Các chức năng chính", 2)
    add_table(doc, ["Chức năng", "Mô tả"], MAIN_FUNCTION_ROWS, [2.2, 5.8], first_col_bold=True)
    add_heading(doc, "4. Use Case Tổng Quát", 2)
    add_image(doc, "use-case.png", "Hình 2.1. Use Case tổng quát của hệ thống", width=6.6)
    add_table(doc, ["Mã", "Use case", "Tác nhân", "Mô tả ngắn"], USE_CASE_ROWS, [0.9, 2.3, 1.5, 3.3], first_col_bold=True, font_size=10.3)
    add_heading(doc, "5. Workflow đặt sân", 2)
    add_body(doc, "Workflow đặt sân là luồng nghiệp vụ trung tâm của hệ thống, bắt đầu từ khai báo tài nguyên, tiếp nhận yêu cầu đặt sân, kiểm tra ràng buộc, phê duyệt booking và kết thúc ở usage log và báo cáo.")
    add_image(doc, "workflow.png", "Hình 2.2. Workflow đặt sân từ khai báo tài nguyên đến usage log", width=6.8)
    add_numbered(doc, WORKFLOW_STEPS)
    add_heading(doc, "6. Yêu cầu chức năng", 2)
    add_bullets(doc, FUNCTIONAL_REQUIREMENTS)
    add_note(doc, "Ràng buộc quan trọng: sân bảo trì không được đặt; booking không được trùng sân - trùng ngày - trùng ca; sinh viên chỉ được thao tác trên booking của chính mình.")
    add_heading(doc, "7. Yêu cầu phi chức năng", 2)
    add_bullets(doc, NON_FUNCTIONAL_REQUIREMENTS)

    add_heading(doc, "CHƯƠNG III: Thiết kế hệ thống", 1)
    add_heading(doc, "1. Sơ đồ kiến trúc MVC", 2)
    add_body(doc, "Dự án được xây dựng trên Laravel theo mô hình MVC. Request từ trình duyệt đi qua route, middleware và controller; controller thực hiện validation, thao tác model và trả kết quả về Blade View. Cách tổ chức này giúp mã nguồn dễ đọc, dễ bảo trì và phù hợp với yêu cầu môn học.")
    add_image(doc, "mvc.png", "Hình 3.1. Kiến trúc MVC của dự án Laravel", width=6.7)
    add_heading(doc, "2. Sơ đồ module", 2)
    add_body(doc, "Hệ thống được tách thành các module theo nhóm nghiệp vụ chính, giúp việc phân chia chức năng, mã nguồn và dữ liệu trở nên rõ ràng hơn.")
    add_image(doc, "module-diagram.png", "Hình 3.2. Sơ đồ module của hệ thống", width=6.8)
    add_table(doc, ["Module", "Nội dung"], MODULE_ROWS, [2.1, 5.9], first_col_bold=True)
    add_heading(doc, "3. Biểu đồ ERD", 2)
    add_image(doc, "erd-overview.png", "Hình 3.3. Biểu đồ ERD tổng quan của hệ thống", width=6.9)
    add_heading(doc, "4. Thiết kế CSDL", 2)
    add_table(doc, ["Nhóm bảng", "Các bảng", "Ý nghĩa"], TABLE_GROUPS, [1.8, 2.8, 3.4], first_col_bold=True)
    add_note(doc, "Ngoài 8 bảng nghiệp vụ cốt lõi, trong phpMyAdmin có thể xuất hiện thêm các bảng như migrations, cache, jobs, sessions. Đây là các bảng hạ tầng của Laravel, không phải bảng nghiệp vụ chính của đề tài.")
    add_table(doc, ["Bảng nguồn", "Bảng đích", "Kiểu", "Ý nghĩa"], RELATIONSHIP_ROWS, [1.5, 2.3, 0.9, 3.1], first_col_bold=True, font_size=10.2)
    add_heading(doc, "5. Data dictionary", 2)
    add_body(doc, "Phần này tổng hợp các cột quan trọng, kiểu dữ liệu và ràng buộc chính của từng bảng nghiệp vụ. Nội dung này giúp minh chứng rõ cho phần thiết kế dữ liệu và quan hệ trong hệ thống.")
    for idx, (name, note, rows) in enumerate(DATA_DICTIONARY, start=1):
        add_heading(doc, f"5.{idx}. Bảng {name}", 3)
        add_body(doc, note)
        add_table(doc, ["Cột", "Kiểu dữ liệu", "Ràng buộc", "Ý nghĩa"], rows, [1.7, 1.8, 2.2, 2.7], first_col_bold=True, font_size=10.0)
    add_heading(doc, "6. Ma trận phân quyền", 2)
    add_table(doc, ["Chức năng", "Admin", "Sinh viên", "Ghi chú"], PERMISSION_ROWS, [3.0, 1.0, 1.3, 2.9], first_col_bold=True, font_size=10.0)

    add_heading(doc, "CHƯƠNG IV: Xây dựng và cài đặt hệ thống", 1)
    add_heading(doc, "1. Công nghệ sử dụng", 2)
    add_table(doc, ["Công nghệ", "Vai trò", "Lý do lựa chọn"], TECHNOLOGY_ROWS, [1.7, 2.2, 4.1], first_col_bold=True)
    add_heading(doc, "2. Môi trường cài đặt", 2)
    add_table(doc, ["Thành phần", "Giá trị", "Vai trò"], ENVIRONMENT_ROWS, [1.8, 1.8, 4.4], first_col_bold=True, font_size=10.3)
    add_heading(doc, "3. Các chức năng đã xây dựng", 2)
    add_table(doc, ["Nhóm chức năng", "Nội dung đã xây dựng"], COMPLETED_FEATURE_ROWS, [2.4, 5.6], first_col_bold=True)
    add_note(doc, "Tài khoản demo hiện có: admin@campus.local / password, student1@campus.local / password, student2@campus.local / password.")
    add_heading(doc, "4. Minh họa giao diện", 2)
    add_body(doc, "Dưới đây là hình minh họa cho ba màn hình đại diện của hệ thống gồm trang đăng nhập, dashboard quản trị và form đặt sân.")
    add_image(doc, "ui-overview.png", "Hình 4.1. Minh họa một số màn hình chính của hệ thống", width=6.9)
    add_note(doc, "Khi demo trực tiếp, có thể đăng nhập bằng tài khoản admin để minh họa dashboard, CRUD danh mục và phê duyệt booking; sau đó đăng nhập bằng tài khoản sinh viên để minh họa form đặt sân.")

    add_heading(doc, "Kết luận", 1)
    add_body(doc, "Đề tài “Xây dựng hệ thống quản lý và đặt sân thể thao tại TLU” đã đạt được mục tiêu chính trong phạm vi vòng 1: hình thành một hệ thống web có quy trình rõ ràng, có kiểm soát xung đột, có phân quyền người dùng và có khả năng báo cáo sử dụng sân. Kết quả này tạo nên một nền tảng tốt để tiếp tục mở rộng thêm các tính năng nâng cao trong các vòng sau, đồng thời đáp ứng phần lớn tiêu chí điểm nền của môn học.")
    add_heading(doc, "Đánh giá tổng quát", 2)
    add_note(doc, "Mức độ hoàn thành hiện tại có thể xem là đã đáp ứng tốt nhóm tiêu chí A1 đến A8. Hai nội dung cần hoàn thiện thêm trước khi nộp là bổ sung thông tin cá nhân trên bìa, chốt commit minh chứng theo mốc và rà soát README / ERD để đồng bộ với bản nộp cuối.")
    add_heading(doc, "Hạn chế hiện tại", 2)
    add_bullets(doc, LIMITATIONS)
    add_heading(doc, "Hướng phát triển", 2)
    add_bullets(doc, FUTURE_DIRECTIONS)
    return doc


def main() -> None:
    doc = build_report()
    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
